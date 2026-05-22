from docx import Document
from docx.shared import Pt, Cm, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


def set_font(run, size=14, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold


def create_base_doc():
    document = Document()

    section = document.sections[0]
    section.left_margin = Mm(30)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    normal = document.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return document


def _centered(document, text, bold=False, size=14):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold)
    return p


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_title_page(document, report, student_name=""):
    year = datetime.now().year

    _centered(document, "Міністерство освіти і науки України")
    _centered(document, "Харківський національний університет радіоелектроніки")
    _centered(document, "")
    _centered(document, "Кафедра штучного інтелекту")

    for _ in range(4):
        _centered(document, "")

    _centered(document, f"Дисципліна «{report.discipline}»")
    _centered(document, report.title)

    for _ in range(5):
        _centered(document, "")

    # Виконав / Прийняла side-by-side with borderless table
    table = document.add_table(rows=3, cols=2)
    _remove_table_borders(table)

    rows_data = [
        ("Виконав:", "Прийняла:"),
        (f"ст. гр. {report.group}", report.teacher),
        (student_name, ""),
    ]
    for row_idx, (left, right) in enumerate(rows_data):
        for col_idx, text in enumerate((left, right)):
            cell = table.cell(row_idx, col_idx)
            para = cell.paragraphs[0]
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(text)
            set_font(run)

    for _ in range(5):
        _centered(document, "")

    _centered(document, f"Харків – {year}")

    document.add_page_break()


def add_table_of_contents(document, headings=None):
    _centered(document, "ЗМІСТ", bold=True)

    # Field begin + instrText + separate — Word uses this to auto-rebuild the TOC
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run()

    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), 'begin')
    r._r.append(fc)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r._r.append(instr)

    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'separate')
    r._r.append(fc2)

    # Pre-populated entries visible before Word updates the field
    for text, level in (headings or []):
        ep = document.add_paragraph()
        ep.paragraph_format.first_line_indent = Cm(0)
        ep.paragraph_format.space_before = Pt(0)
        ep.paragraph_format.space_after = Pt(2)
        ep.paragraph_format.left_indent = Cm((level - 1) * 1.25)
        set_font(ep.add_run(text), size=14)

    # Field end
    p_end = document.add_paragraph()
    p_end.paragraph_format.first_line_indent = Cm(0)
    p_end.paragraph_format.space_before = Pt(0)
    p_end.paragraph_format.space_after = Pt(0)
    r_end = p_end.add_run()
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'end')
    r_end._r.append(fc3)

    document.add_page_break()


def add_heading(document, text, level=1):
    """Use Word built-in heading styles so the TOC field picks them up."""
    paragraph = document.add_paragraph(style=f'Heading {level}')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        run = paragraph.add_run(text.upper())
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
        run = paragraph.add_run(text)

    set_font(run, bold=True)
    return paragraph

def add_paragraph_text(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_font(run)
    return paragraph    

def add_listing(document, code, listing_number, title):
    caption = document.add_paragraph()
    caption_run = caption.add_run(f"Лістинг {listing_number} – {title}")
    set_font(caption_run)
    caption.paragraph_format.first_line_indent = Cm(0)
    
    code_paragraph = document.add_paragraph()
    code_run = code_paragraph.add_run(code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(12)
    code_paragraph.paragraph_format.first_line_indent = Cm(0)
    code_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return code_paragraph


def add_image(document, image_path, image_number, caption_text):
    img_paragraph = document.add_paragraph()
    img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_paragraph.paragraph_format.first_line_indent = Cm(0)
    run = img_paragraph.add_run()
    run.add_picture(image_path, width=Inches(5))

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption_run = caption.add_run(f"Рисунок {image_number} – {caption_text}")
    set_font(caption_run)
    return caption

def build_report_docx(report, images=None, student_name=""):
    document = create_base_doc()
    add_title_page(document, report, student_name)

    toc_headings = [("Мета роботи", 1)]
    if report.sections:
        for idx, sec in enumerate(report.sections):
            toc_headings.append((sec.get('title', f'Розділ {idx + 1}'), 1))
    toc_headings.append(("Висновки", 1))
    add_table_of_contents(document, toc_headings)

    add_heading(document, "Мета роботи")
    add_paragraph_text(document, report.goal or "")

    if report.sections:
        images_by_section = {}
        if images:
            for img in images:
                images_by_section.setdefault(img.section_path, []).append(img)

        listing_counter = 1
        image_counter = 1
        for idx, section in enumerate(report.sections):
            title = section.get('title', f'Розділ {idx + 1}')
            add_heading(document, title, level=1)
            if section.get('text'):
                add_paragraph_text(document, section['text'])
            if section.get('code'):
                add_listing(document, section['code'], listing_counter, title)
                listing_counter += 1
            for img in images_by_section.get(str(idx), []):
                img_path = os.path.join("/app/uploads", img.filename)
                if os.path.exists(img_path):
                    add_image(document, img_path, image_counter, img.caption)
                    image_counter += 1

    add_heading(document, "Висновки")
    add_paragraph_text(document, report.conclusion or "")

    output_dir = "/app/generated"
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"report_{report.id}.docx")
    document.save(path)
    return path
